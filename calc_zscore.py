import datetime
import os

import matplotlib
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from scipy import signal
from scipy.integrate import simps

from topograph import get_psds_alpha, get_psds_beta, get_psds_theta, plot_topomap

matplotlib.use("Agg")


# Calculating alpha power difference between Fp1 and Fp2.
def calc_alpha_diff(alpha_fp1, alpha_fp2):
    calc_result = np.log(alpha_fp2) - np.log(alpha_fp1)
    return calc_result


def calc_zscore(input_data, input_name):
    input_data = os.path.join(os.getcwd(), "out", input_data)
    out_dir = os.path.join(os.getcwd(), "out")

    df = pd.read_table(
        input_data,
        header=None,
        delim_whitespace=True,
        names=(
            "Date1",
            "Date2",
            "Fp1",
            "Fp2",
            "F7",
            "F8",
            "C3",
            "C4",
            "T3",
            "T4",
            "T5",
            "T6",
            "O1",
            "O2",
            "X1",
            "Other1",
            "Other2",
        ),
    )

    # dfの列を並び替える
    df = df[
        [
            "Date1",
            "Date2",
            "Fp1",
            "Fp2",
            "C3",
            "C4",
            "O1",
            "O2",
            "T3",
            "T4",
            "X1",
            "F7",
            "F8",
            "T5",
            "T6",
            "Other1",
            "Other2",
        ]
    ]

    df.drop(columns=["Date1", "Date2", "X1", "Other1", "Other2"], inplace=True)
    # 欠損値が１つでもある行を削除する
    df = df.dropna(how="any")

    # topomapを作成して保存する
    df_copy = df.T
    df_np = df_copy.values

    # シータ波のスペクトルを求める
    pwrs_theta, _ = get_psds_theta(df_np)

    # アルファ波のスペクトルを求める
    pwrs_alpha, _ = get_psds_alpha(df_np)

    # ベータ波のスペクトルを求める
    pwrs_beta, _ = get_psds_beta(df_np)

    # Zスコアを求める
    eeg_list = [
        "Fp1",
        "Fp2",
        "C3",
        "C4",
        "O1",
        "O2",
        "T3",
        "T4",
        "F7",
        "F8",
        "T5",
        "T6",
    ]

    npy_dir = os.path.join(os.getcwd(), "npy")
    np_load_dataset = [
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_fp1_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_fp1_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_fp1_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_fp2_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_fp2_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_fp2_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_c3_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_c3_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_c3_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_c4_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_c4_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_c4_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_O1_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_O1_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_O1_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_O2_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_O2_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_O2_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_t3_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_t3_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_t3_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_t4_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_t4_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_t4_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_f7_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_f7_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_f7_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_f8_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_f8_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_f8_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_t5_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_t5_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_t5_eval+train.npy"),
        ],
        [
            os.path.join(npy_dir, "pwrs_rel_alpha_t6_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_beta_t6_eval+train.npy"),
            os.path.join(npy_dir, "pwrs_rel_theta_t6_eval+train.npy"),
        ],
    ]

    j_list = ["アルファ", "ベータ", "シータ"]

    # result_listは結果を文章で出力するために用意したリスト
    result_list = []

    # ３行１６列のリストを作る。そしてアルファ波のtopomapならlist[0][l]をlでイテレートして表示する。numpy配列にしないといけないはず。
    # →そのやり方が分からなかったので３種類のリストを作る！
    result_alpha = []
    result_beta = []
    result_theta = []

    sample_spectrum = []

    alpha_fp1, alpha_fp2 = 0, 0
    rel_alpha_fp1, rel_alpha_fp2 = 0, 0
    theta_beta_fp1, theta_beta_fp2 = 0, 0

    # ある電極での相対スペクトル密度のZ値を求める
    for i in range(len(eeg_list)):
        sf = 500.0
        df_analyze = df.iloc[:, i]
        df_analyze_np = df_analyze.values

        # Welch's periodogramを求める
        win = 4 * sf
        freqs, psd = signal.welch(df_analyze_np, sf, nperseg=win)
        freq_res = freqs[1] - freqs[0]

        # ベータ波の相対スペクトル密度を求める
        idx_beta = np.logical_and(freqs >= 12, freqs <= 30)
        beta_power = simps(psd[idx_beta], dx=freq_res)

        # アルファ波の相対スペクトル密度を求める
        idx_alpha = np.logical_and(freqs >= 8, freqs <= 12)
        alpha_power = simps(psd[idx_alpha], dx=freq_res)

        # シータ波の相対スペクトル密度を求める
        idx_theta = np.logical_and(freqs >= 4, freqs <= 8)
        theta_power = simps(psd[idx_theta], dx=freq_res)

        total_power = theta_power + alpha_power + beta_power
        relative_alpha = alpha_power / total_power
        relative_beta = beta_power / total_power
        relative_theta = theta_power / total_power
        relative_list = [relative_alpha, relative_beta, relative_theta]

        # 脳波データの周波数スペクトルを求める
        x = freqs
        y = psd
        plt.plot(x, y)
        plt.title("EEG-" + eeg_list[i])
        plt.xlabel("Frequency [Hz]")
        plt.ylabel("Power [μV]")
        plt.xlim(0, 20)
        plt.savefig(os.path.join(out_dir, "EEG-" + eeg_list[i] + ".png"))
        plt.clf()

        # それぞれの電極で計算する

        for j in range(3):
            # 標準偏差を求める
            sample_spectrum = np.load(np_load_dataset[i][j])
            data_mean = np.mean(sample_spectrum)
            data_std = np.std(sample_spectrum)

            z = (relative_list[j] - data_mean) / data_std
            result_list.append(
                "{}電極の{}波の相対パワースペクトルのZ値は{}です。".format(
                    eeg_list[i], j_list[j], round(z, 2)
                )
            )

        z1 = (relative_list[0] - data_mean) / data_std
        result_alpha.append(z1)

        z2 = (relative_list[1] - data_mean) / data_std
        result_beta.append(z2)

        z3 = (relative_list[2] - data_mean) / data_std
        result_theta.append(z3)

        # Fp1とFp2のアルファ波パワーの差とZ値を求める
        # Fp1とFp2のシータ波パワー/ベータ波パワーとZ値を求める
        if eeg_list[i] == "Fp1":
            alpha_fp1 = alpha_power
            rel_alpha_fp1 = relative_alpha
            theta_beta_fp1 = theta_power / beta_power  # = rel_theta / rel_beta

            # theta_betaの標準偏差を求める
            spectrum_theta_beta_fp1 = np.load(np_load_dataset[0][2]) / np.load(
                np_load_dataset[0][1]
            )
            z_theta_beta_fp1 = (
                theta_beta_fp1 - np.mean(spectrum_theta_beta_fp1)
            ) / np.std(spectrum_theta_beta_fp1)

        elif eeg_list[i] == "Fp2":
            alpha_fp2 = alpha_power
            rel_alpha_fp2 = relative_alpha
            theta_beta_fp2 = theta_power / beta_power  # = rel_theta / rel_beta

            # rel_alpha_diffの標準偏差を求める
            rel_alpha_diff = np.log(rel_alpha_fp2) - np.log(rel_alpha_fp1)
            sample_spectrum_diff = np.log(np.load(np_load_dataset[1][0])) - np.log(
                np.load(np_load_dataset[0][0])
            )
            z_alpha = (rel_alpha_diff - np.mean(sample_spectrum_diff)) / np.std(
                sample_spectrum_diff
            )

            # Calculate the difference between the alpha power Fp1 with Fp2
            calc_result = calc_alpha_diff(alpha_fp1, alpha_fp2)

            # theta_betaの標準偏差を求める
            spectrum_theta_beta_fp2 = np.load(np_load_dataset[1][2]) / np.load(
                np_load_dataset[1][1]
            )
            z_theta_beta_fp2 = (
                theta_beta_fp2 - np.mean(spectrum_theta_beta_fp2)
            ) / np.std(spectrum_theta_beta_fp2)

    # シータ波のtopomapを出力する
    fig, ax = plt.subplots(figsize=(10, 8))
    plot_topomap(result_theta, ax, fig)
    plt.title("theta_topomap")
    plt.savefig(os.path.join(out_dir, "theta_save_topomap.png"))

    fig, ax = plt.subplots(figsize=(10, 8))
    plot_topomap(result_alpha, ax, fig)
    plt.title("alpha_topomap")
    plt.savefig(os.path.join(out_dir, "alpha_save_topomap.png"))

    fig, ax = plt.subplots(figsize=(10, 8))
    plot_topomap(result_beta, ax, fig)
    plt.title("beta_topomap")
    plt.savefig(os.path.join(out_dir, "beta_save_topomap.png"))

    plt.clf()

    dt_now = datetime.datetime.now()
    document = Document()
    document.add_heading("結果報告書")
    document.add_paragraph(" ")
    document.add_paragraph("作成日：" + dt_now.strftime("%Y年%m月%d日"))

    file_name = os.path.splitext(os.path.basename(input_data))[0] + ".m00"
    document.add_paragraph("ファイル名：" + file_name)
    document.add_paragraph("インプット名：" + input_name)
    document.add_paragraph(" ")
    document.add_paragraph("あなたのアルファ波の分布は以下のようになります。")
    document.add_picture(
        os.path.join(out_dir, "alpha_save_topomap.png"), width=Inches(3.5)
    )
    document.add_paragraph("あなたのベータ波の分布は以下のようになります。")
    document.add_picture(
        os.path.join(out_dir, "beta_save_topomap.png"), width=Inches(3.5)
    )
    document.add_paragraph("あなたのシータ波の分布は以下のようになります。")
    document.add_picture(
        os.path.join(out_dir, "theta_save_topomap.png"), width=Inches(3.5)
    )

    document.add_paragraph(
        "あなたの脳波検査の結果からZスコアを算出します。（Zスコアは1297人の正常被験者のデータセットから算出しております）"
    )

    for k in range(len(result_list)):
        document.add_paragraph(" ")
        document.paragraphs[13 + k].add_run(result_list[k])

    document.add_paragraph(
        "前頭葉のアルファ波左右差の値は" + str("{:.3g}".format(calc_result)) + "です。"
    )
    document.add_paragraph(
        "前頭葉のアルファ波左右差の相対パワースペクトルのZ値は" + str("{:.3g}".format(z_alpha)) + "です。"
    )
    document.add_paragraph(
        "前頭葉（右）のシータ波/ベータ波の比率は" + str("{:.3g}".format(theta_beta_fp1)) + "です。"
    )
    document.add_paragraph(
        "前頭葉（左）のシータ波/ベータ波の比率は" + str("{:.3g}".format(theta_beta_fp2)) + "です。"
    )
    document.add_paragraph(
        "前頭葉（右）のシータ波/ベータ波のZ値は" + str("{:.3g}".format(z_theta_beta_fp1)) + "です。"
    )
    document.add_paragraph(
        "前頭葉（左）のシータ波/ベータ波のZ値は" + str("{:.3g}".format(z_theta_beta_fp2)) + "です。"
    )
    for i in range(len(eeg_list)):
        document.add_picture(
            os.path.join(out_dir, "EEG-" + eeg_list[i] + ".png"), width=Inches(3.5)
        )

    save_dir = input_data.replace(".txt", ".docx")

    document.save(save_dir)
